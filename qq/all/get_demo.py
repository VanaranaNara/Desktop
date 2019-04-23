# !/usr/bin/python
# -*- coding:utf-8 -*-
# @Time      : 2019/4/8 14:19
# @Author    : Alice
# @Site      : 
# @File      : get_demo.py
# @Software  : PyCharm
import requests
import os
import re
from bs4 import BeautifulSoup
from urllib import request
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class GetCsdn(object):
    def __init__(self, url):
        self.url = url
        self.doc = Document()
        self.name = ''
        self.path = ''

    def get_html(self, url):
        header = self.get_header()
        req = requests.get(url, headers=header)
        req = req.content.decode('utf-8')  # 转换文件的编码格式
        return req

    def get_header(self):
        header = {
            'Referer': 'https://blog.csdn.net/Ly4wU5giY/article/details/7971055',
            'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.25 Safari/537.36 Core/1.70.3641.400 QQBrowser/10.4.3284.400'
        }
        return header

    def get_title(self, soup):
        """
        获取文章标题
        :param soup: soup 对像
        :return:
        """
        com = re.compile(r'e>(.*?)</', re.S)
        title = soup.select('title')
        title = re.findall(com, str(title))[0]
        if not os.path.exists(title):
            os.mkdir(title)
        self.path = title + '/'
        self.name = title + '/' + title

        # print(title)

    def write_p(self, strs):
        """
        向文本里添加段落
        :param strs:
        :return:
        """
        paragraph = self.doc.add_paragraph(strs)
        ph_format = paragraph.paragraph_format
        ph_format.line_spacing = Pt(14)  # 设置行间距

    def write_img(self, imgs, num):
        img1 = re.compile(r'="(.*?)"', re.S)
        img = re.findall(img1, str(imgs))
        save_path1 = self.path + str(num) + '.jpeg'
        request.urlretrieve(img[2], save_path1)  # 下载请求
        self.doc.add_picture(save_path1)
        last_paragraph = self.doc.paragraphs[-1]  # 段落属性，在这里代表每一行，一共三行，-1为最后一行
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置

    def run(self):
        html = self.get_html(self.url)
        soup = BeautifulSoup(html, 'html.parser')
        self.get_title(soup)  # 获取文章标题
        print(self.path)
        p = soup.select('#content_views #js_content p')
        if not p:
            p = soup.select('#content_views p')
        num = 1
        for i in p:
            imgs = i.find('img')
            if imgs:
                self.write_img(imgs, num)
                num += 1
            else:
                urls2 = i.get_text()
                if urls2 == ' ':
                    continue
                else:
                    pass
                    self.write_p(urls2)
        self.doc.save(self.name + '.docx')


# my_url = 'https://blog.csdn.net/Ly4wU5giY/article/details/79710559'  # 待爬取页面
# my_url = 'https://blog.csdn.net/chengxuyuan997/article/details/81231879'  # 待爬取页面

while True:
    my_url = input("请输入你要爬取的网址：")
    if my_url != '':
        demo = GetCsdn(my_url)
        demo.run()
